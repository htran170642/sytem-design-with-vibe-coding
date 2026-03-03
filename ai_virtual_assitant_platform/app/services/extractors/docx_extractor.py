"""
DOCX Text Extractor
Extract text and metadata from DOCX files
Phase 4, Step 2: Document parsing and text chunking
"""

from pathlib import Path
from typing import Tuple, Dict, Any
from docx import Document

from app.core.logging_config import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_path: Path) -> Tuple[str, Dict[str, Any]]:
    """
    Extract text and metadata from DOCX file
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (text_content, metadata)
        
    Example:
        >>> text, metadata = extract_text_from_docx("document.docx")
        >>> print(metadata["paragraphs"])
        50
    """
    try:
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        full_text = "\n\n".join(text_parts)
        
        # Extract metadata
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "created_date": str(doc.core_properties.created) if doc.core_properties.created else None,
            "word_count": len(full_text.split()),
        }
        
        logger.info(
            "DOCX extracted",
            extra={
                "file": file_path.name,
                "paragraphs": metadata["paragraphs"],
                "word_count": metadata["word_count"],
            }
        )
        
        return full_text, metadata
        
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}", exc_info=True)
        raise Exception(f"Failed to extract DOCX: {str(e)}")